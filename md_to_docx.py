"""
Markdown 转 Word 文档工具（支持 LaTeX 公式）
用法: python md_to_docx.py input.md [output.docx]
"""

import re
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from latex_utils import simplify_tex

# ==================== 字体全局配置 ====================
BODY_FONT = '微软雅黑'
BODY_SIZE = Pt(11)
MATH_FONT = 'Cambria Math'
MATH_SIZE = Pt(11)
MATH_COLOR = RGBColor(0x33, 0x33, 0x33)
CODE_FONT = 'Consolas'
CODE_SIZE = Pt(9.5)
HEADING_COLOR = RGBColor(0x1A, 0x56, 0xDB)


def _ensure_font(rPr, name):
    """确保 run 的 rFonts 中西文字体一致"""
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {qn("w:eastAsia")}="{name}" '
            f'{qn("w:ascii")}="{name}" {qn("w:hAnsi")}="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), name)
        rFonts.set(qn('w:ascii'), name)
        rFonts.set(qn('w:hAnsi'), name)


def _make_run(paragraph, text='', font_name=BODY_FONT, font_size=BODY_SIZE,
              bold=None, italic=None, color=None):
    """创建格式统一的 run，确保中英文字体一致"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    _ensure_font(run._element.get_or_add_rPr(), font_name)
    return run


def setup_styles(doc: Document):
    """配置文档全局样式"""
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    _ensure_font(style.element.get_or_add_rPr(), BODY_FONT)

    for lv in range(1, 4):
        hs = doc.styles[f'Heading {lv}']
        hs.font.color.rgb = HEADING_COLOR
        hs.font.name = BODY_FONT
        _ensure_font(hs.element.get_or_add_rPr(), BODY_FONT)

    for sn in ['List Bullet', 'List Number']:
        try:
            ls = doc.styles[sn]
            ls.font.name = BODY_FONT
            ls.font.size = BODY_SIZE
            _ensure_font(ls.element.get_or_add_rPr(), BODY_FONT)
        except KeyError:
            pass


# ==================== LaTeX 简化（使用共享模块 latex_utils） ====================
# simplify_tex 从 latex_utils 导入，见文件顶部


# ==================== Markdown 解析 ====================

def _merge_display_math(lines: list) -> list:
    """跨行 $$...$$ 合并为单行"""
    result = []
    in_math = False
    buf = []
    for line in lines:
        s = line.strip()
        if s.startswith('$$') and not in_math:
            if s.endswith('$$') and len(s) > 4:
                result.append(s)
                continue
            in_math = True
            buf = [s]
        elif s.endswith('$$') and in_math:
            buf.append(s)
            in_math = False
            merged = '$$' + '\n'.join(l.strip('$') for l in buf).strip() + '$$'
            result.append(merged)
        elif in_math:
            buf.append(s)
        else:
            result.append(line)
    return result


def parse_markdown(lines: list) -> list:
    """解析 Markdown → 结构化元素"""
    merged = _merge_display_math(lines)
    elements = []
    i = 0
    while i < len(merged):
        line = merged[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # 块级公式 $$...$$
        if s.startswith('$$') and s.endswith('$$') and len(s) > 4:
            elements.append({'type': 'display_math', 'tex': s[2:-2].strip()})
            i += 1
            continue

        # 代码块 ```
        if s.startswith('```'):
            lang = s[3:].strip()
            code_lines = []
            i += 1
            while i < len(merged) and not merged[i].strip().startswith('```'):
                code_lines.append(merged[i])
                i += 1
            i += 1
            elements.append({'type': 'code_block', 'code': '\n'.join(code_lines), 'lang': lang})
            continue

        # 水平线 ---
        if re.match(r'^[-*_]{3,}\s*$', s):
            elements.append({'type': 'hr'})
            i += 1
            continue

        # 标题 #
        hm = re.match(r'^(#{1,6})\s+(.+)', s)
        if hm:
            elements.append({'type': 'heading', 'level': len(hm.group(1)), 'text': hm.group(2)})
            i += 1
            continue

        # 无序列表 - * +
        lm = re.match(r'^(\s*)[-*+]\s+(.+)', s)
        if lm:
            items = []
            # 关键修复：用原始 merged[i] 计算 indent，而非 strip 后的 s
            indent = len(re.match(r'^(\s*)', merged[i]).group(1))
            while i < len(merged):
                cm = re.match(r'^(\s*)[-*+]\s+(.+)', merged[i])
                if cm and len(cm.group(1)) == indent:
                    items.append(cm.group(2))
                    i += 1
                elif merged[i].strip() == '':
                    i += 1
                    if i < len(merged):
                        nm = re.match(r'^(\s*)[-*+]\s+(.+)', merged[i])
                        if nm and len(nm.group(1)) == indent:
                            items.append(nm.group(2))
                            i += 1
                            continue
                        i -= 1
                    break
                else:
                    break
            elements.append({'type': 'ul', 'items': items})
            continue

        # 有序列表 1.
        om = re.match(r'^(\s*)\d+\.\s+(.+)', s)
        if om:
            items = []
            while i < len(merged):
                dm = re.match(r'^(\s*)\d+\.\s+(.+)', merged[i])
                if dm:
                    items.append(dm.group(2))
                    i += 1
                elif merged[i].strip() == '':
                    i += 1
                    if i < len(merged) and re.match(r'^(\s*)\d+\.\s+(.+)', merged[i]):
                        continue
                    i -= 1
                    break
                else:
                    break
            elements.append({'type': 'ol', 'items': items})
            continue

        # 引用块 >
        qm = re.match(r'^>\s?(.*)', s)
        if qm:
            qlines = []
            while i < len(merged):
                mm = re.match(r'^>\s?(.*)', merged[i])
                if mm:
                    qlines.append(mm.group(1))
                    i += 1
                elif merged[i].strip() == '':
                    i += 1
                    if i < len(merged) and re.match(r'^>\s?(.*)', merged[i]):
                        continue
                    i -= 1
                    break
                else:
                    break
            elements.append({'type': 'blockquote', 'text': '\n'.join(qlines)})
            continue

        # 表格 |
        if '|' in s and i + 1 < len(merged) and re.match(r'^[\|\-:\s]+$', merged[i + 1].strip()):
            tlines = [s]
            i += 2
            while i < len(merged) and '|' in merged[i]:
                tlines.append(merged[i].strip())
                i += 1
            rows = [[c.strip() for c in tl.strip('|').split('|')] for tl in tlines]
            if rows:
                elements.append({'type': 'table', 'rows': rows})
            continue

        # 普通段落
        plines = [s]
        i += 1
        while i < len(merged) and merged[i].strip() and \
              not re.match(r'^(#{1,6}\s|```|[-*_]{3,}|[-*+]\s|\d+\.\s|>\s|\$\$)',
                           merged[i].strip()) and '|' not in merged[i]:
            plines.append(merged[i].strip())
            i += 1
        elements.append({'type': 'paragraph', 'text': ' '.join(plines)})

    return elements


# ==================== 行内格式处理 ====================

def _add_math_text(paragraph, tex: str, inline: bool = False):
    """将 LaTeX 公式以数学字体格式插入"""
    clean = simplify_tex(tex)
    if inline:
        _make_run(paragraph, clean, font_name=MATH_FONT, font_size=Pt(10.5),
                  italic=True, color=MATH_COLOR)
    else:
        _make_run(paragraph, clean, font_name=MATH_FONT, font_size=MATH_SIZE,
                  italic=True, color=MATH_COLOR)


def _add_hyperlink(paragraph, text: str, url: str):
    """在 Word 段落中添加可点击的超链接"""
    from docx.oxml.ns import qn
    from lxml import etree
    # 创建超链接关系
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    # 手动构建 XML 元素避免命名空间问题
    NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

    hyperlink = etree.SubElement(paragraph._element, f'{{{NS_W}}}hyperlink')
    hyperlink.set(f'{{{NS_R}}}id', r_id)

    run = etree.SubElement(hyperlink, f'{{{NS_W}}}r')
    rpr = etree.SubElement(run, f'{{{NS_W}}}rPr')
    # 超链接样式
    rstyle = etree.SubElement(rpr, f'{{{NS_W}}}rStyle')
    rstyle.set(f'{{{NS_W}}}val', 'Hyperlink')
    # 字体颜色
    color = etree.SubElement(rpr, f'{{{NS_W}}}color')
    color.set(f'{{{NS_W}}}val', '1A56DB')
    # 下划线
    u = etree.SubElement(rpr, f'{{{NS_W}}}u')
    u.set(f'{{{NS_W}}}val', 'single')
    # 字体
    rfonts = etree.SubElement(rpr, f'{{{NS_W}}}rFonts')
    rfonts.set(f'{{{NS_W}}}eastAsia', BODY_FONT)
    rfonts.set(f'{{{NS_W}}}ascii', BODY_FONT)
    rfonts.set(f'{{{NS_W}}}hAnsi', BODY_FONT)

    t = etree.SubElement(run, f'{{{NS_W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text


def process_inline_text(paragraph, text: str):
    """处理行内格式：$...$ 公式、**粗体**、*斜体*、`代码`、[链接]"""
    math_items = []

    def _save_math(m):
        math_items.append(m.group(1))
        return f'[MATH{len(math_items) - 1}]'

    # 先保护 $ 公式
    text = re.sub(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', _save_math, text)

    # 注意：模式顺序很重要，长模式在前
    pattern = (
        r'(\*\*\*(.+?)\*\*\*|'      # ***bold italic***
        r'___\s*(.+?)\s*___|'          # ___bold italic___
        r'\*\*(.+?)\*\*|'              # **bold**
        r'__\s*(.+?)\s*__|'             # __bold__
        r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|'  # *italic* (不匹配 **)
        r'(?<!_)_(?!_)\s*(.+?)\s*(?<!_)_(?!_)|'      # _italic_
        r'`(.+?)`|'                     # `code`
        r'!\[(.*?)\]\((.+?)\)|'      # ![alt](url)
        r'\[(.+?)\]\((.+?)\))'        # [text](url)
    )

    last = 0
    for m in re.finditer(pattern, text):
        start, end = m.span()
        if start > last:
            _emit_segment(paragraph, text[last:start], math_items)
        last = end

        g = m.group(0)
        if m.group(2) and g.startswith('***'):
            _make_run(paragraph, m.group(2), bold=True, italic=True)
        elif m.group(3) and g.startswith('___'):
            _make_run(paragraph, m.group(3), bold=True, italic=True)
        elif m.group(4):
            _make_run(paragraph, m.group(4), bold=True)
        elif m.group(5):
            _make_run(paragraph, m.group(5), bold=True)
        elif m.group(6):
            _make_run(paragraph, m.group(6), italic=True)
        elif m.group(7):
            _make_run(paragraph, m.group(7), italic=True)
        elif m.group(8):
            _make_run(paragraph, m.group(8), font_name=CODE_FONT,
                      font_size=CODE_SIZE, color=RGBColor(0xD0, 0x3A, 0x2E))
        elif m.group(9):
            _make_run(paragraph, f'[图片: {m.group(9)}]',
                      italic=True, color=RGBColor(0x99, 0x99, 0x99))
        elif m.group(11):
            # 可点击超链接
            _add_hyperlink(paragraph, m.group(11), m.group(12))

    if last < len(text):
        _emit_segment(paragraph, text[last:], math_items)


def _emit_segment(paragraph, text: str, math_items: list):
    """输出文本片段，还原公式占位符"""
    parts = re.split(r'(\[MATH(\d+)\])', text)
    i = 0
    while i < len(parts):
        part = parts[i]
        if not part:
            i += 1
            continue
        if part.startswith('[MATH') and i + 1 < len(parts):
            idx = int(parts[i + 1])
            if idx < len(math_items):
                _add_math_text(paragraph, math_items[idx], inline=True)
            i += 2
        else:
            _make_run(paragraph, part)
            i += 1


# ==================== Word 文档写入 ====================

def add_elements_to_doc(doc: Document, elements: list):
    """将解析元素写入 Word"""
    for el in elements:
        t = el['type']

        if t == 'heading':
            h = doc.add_heading(level=min(el['level'], 3))
            h.clear()
            process_inline_text(h, el['text'])
            # 标题所有 run 统一加粗
            for run in h.runs:
                run.bold = True

        elif t == 'display_math':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            pPr = p._element.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'FAFAFA'
            })
            pPr.insert(0, shd)
            _add_math_text(p, el['tex'], inline=False)

        elif t == 'paragraph':
            p = doc.add_paragraph()
            process_inline_text(p, el['text'])

        elif t == 'code_block':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F0F0F0'
            })
            pPr.insert(0, shd)
            _make_run(p, el['code'].rstrip('\n'), font_name=CODE_FONT, font_size=CODE_SIZE)
            if el.get('lang'):
                _make_run(p, f'\n（{el["lang"]}）', font_size=Pt(9),
                          color=RGBColor(0x99, 0x99, 0x99))

        elif t == 'ul':
            for item in el['items']:
                p = doc.add_paragraph(style='List Bullet')
                p.clear()
                process_inline_text(p, item)

        elif t == 'ol':
            for item in el['items']:
                p = doc.add_paragraph(style='List Number')
                p.clear()
                process_inline_text(p, item)

        elif t == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            left = pBdr.makeelement(qn('w:left'), {
                qn('w:val'): 'single', qn('w:sz'): '12',
                qn('w:space'): '8', qn('w:color'): 'CCCCCC'
            })
            pBdr.append(left)
            pPr.insert(0, pBdr)
            _make_run(p, el['text'], italic=True, color=RGBColor(0x66, 0x66, 0x66))

        elif t == 'hr':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single', qn('w:sz'): '6',
                qn('w:space'): '1', qn('w:color'): 'CCCCCC'
            })
            pBdr.append(bottom)
            pPr.insert(0, pBdr)

        elif t == 'table':
            rows = el['rows']
            if not rows:
                continue
            nc = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=nc, style='Light Grid Accent 1')
            for ri, row in enumerate(rows):
                for ci, ct in enumerate(row):
                    if ci < nc:
                        cell = table.cell(ri, ci)
                        cell.text = ''
                        process_inline_text(cell.paragraphs[0], ct)
                        if ri == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True

    doc.add_paragraph()


# ==================== 主流程 ====================

def convert_md_to_docx(input_path: str, output_path: str = None):
    """主函数"""
    if not os.path.exists(input_path):
        print(f"❌ 文件不存在: {input_path}")
        sys.exit(1)

    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.splitlines()
    print(f"📖 已读取: {input_path} ({len(lines)} 行)")

    elements = parse_markdown(lines)
    print(f"🔍 解析出 {len(elements)} 个元素")

    doc = Document()
    setup_styles(doc)

    title = doc.add_heading(os.path.basename(input_path), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_elements_to_doc(doc, elements)

    doc.save(output_path)
    print(f"✅ 已生成: {output_path}")
    return output_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("📝 Markdown → Word 转换工具 (支持 LaTeX 公式)")
        print(f"用法: python {os.path.basename(__file__)} <input.md> [output.docx]")
        print(f"      python {os.path.basename(__file__)} <input_dir> --batch")
        print()
        print("示例:")
        print(f"  python {os.path.basename(__file__)} readme.md")
        print(f"  python {os.path.basename(__file__)} readme.md 输出.docx")
        print(f"  python {os.path.basename(__file__)} ./markdown_files/ --batch")
        sys.exit(0)

    input_path = sys.argv[1]
    if len(sys.argv) > 2 and sys.argv[2] == '--batch':
        import glob
        input_dir = input_path
        if not os.path.isdir(input_dir):
            print(f"❌ 目录不存在: {input_dir}")
            sys.exit(1)
        md_files = glob.glob(os.path.join(input_dir, '*.md'))
        if not md_files:
            print(f"❌ 目录中未找到 .md 文件: {input_dir}")
            sys.exit(1)
        print(f"📦 批量转换 {len(md_files)} 个文件...")
        success = 0
        for f in md_files:
            try:
                convert_md_to_docx(f)
                success += 1
            except Exception as e:
                print(f"❌ 转换失败 {f}: {e}")
        print(f"✅ 批量完成: {success}/{len(md_files)}")
    else:
        convert_md_to_docx(input_path, sys.argv[2] if len(sys.argv) > 2 else None)
